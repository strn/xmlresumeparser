#!/usr/bin/python
# -*- coding: utf-8 -*-


from   docx import Document, enum
from   pprint import pprint
import sys
from   BaseWriter import BaseWriter
import ResumeWriter.Util.Lookup as Lkp



class ResumeWriter(BaseWriter):
	
	"""
	MS Word document writer
	"""

	def __init__(self, filename=None):
		# This number says how many jobs in job history
		# will be listed in full (with all details)
		# The rest will be mentioned just by employer name
		if filename == None:
			print "(E) Can't output ODF file to STDOUT, exiting ..."
			sys.exit(2)
		else:
			self.doc = Document()
		self.filename = filename
		self.jobDivide = 2


	def write(self, model=None, lang='en_GB', withPhoto=False, personalData=False, experience=False):
		if model == None:
			print "(E) Empty resume model."
			return
		self.lang = lang
		self.withPhoto = withPhoto
		self.personalData = personalData
		self.experience = experience
		self.writeHeader(model)
		self.writeAddress(model)
		#self.writeSummary(model)
		#self.writeSkills(model)
		#self.writeEmployment(model)
		#self.writeCertifications(model)
		#self.writeEducation(model)
		#self.writeLanguages(model)
		#self.writeHobbies(model)
		#self.writeReferences(model)
		self.doc.save( self.filename + ".docx" )


	def writeHeader(self, model):
		self.doc.add_heading( model.personName, 0 )


	def writeAddress(self, model):
		if self.withPhoto:
			if model.personPhoto and model.personPhoto[:7] == 'file://':
				# Picture is added using its OS path
				personPhoto = model.personPhoto[7:]
			else:
				personPhoto = model.personPhoto
			# Merge cells in 1st column - place for picture
			if self.personalData:
				table = self.doc.add_table(rows=11, cols=4)
				a = table.cell(0,0)
				b = table.cell(10,0)
			else:
				table = self.doc.add_table(rows=4, cols=3)
				a = table.cell(0,0)
				b = table.cell(3,0)
			a.merge(b)
			# Locate first cell to add picture
			picCell = table.cell(0,0)
			parPicCell = picCell.paragraphs[0]
			picRun = parPicCell.add_run()
			picRun.add_picture( personPhoto )
			adrCol = table.columns[1].cells
		else:
			table = self.doc.add_table(rows=7, cols=2)
			adrCol = table.columns[0].cells
		adrCol[0].text = model.streetAddress
